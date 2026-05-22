from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("TP2 — Sprint 1 (Grupo G4 — WMS)")
run.bold = True
run.font.size = Pt(16)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run("Respostas do README mínimo pedido no enunciado")
sub_run.italic = True
sub_run.font.size = Pt(11)

doc.add_paragraph()


def add_qa(question: str, answer_paragraphs: list[str], bullets: list[str] | None = None):
    q = doc.add_paragraph()
    q_run = q.add_run("Pergunta: ")
    q_run.bold = True
    q.add_run(question)

    for text in answer_paragraphs:
        p = doc.add_paragraph()
        a_run = p.add_run("Resposta: " if text is answer_paragraphs[0] else "")
        a_run.bold = True
        p.add_run(text)

    if bullets:
        for item in bullets:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()


add_qa(
    "Qual arquitetura foi escolhida?",
    [
        "Onion. O domínio fica no centro e não depende de nada. A aplicação depende só do domínio. "
        "A infraestrutura implementa as interfaces declaradas no domínio. A apresentação (API REST) "
        "monta tudo via injeção de dependência.",
        "A regra de dependência é: Presentation → Application → Domain ← Infrastructure.",
        "Escolhemos Onion porque o WMS tem regras de negócio que precisam ficar isoladas do banco e do "
        "framework — principalmente a regra de FIFO da expedição, que dá para testar sem subir HTTP nem banco.",
    ],
)

add_qa(
    "Qual fluxo foi iniciado?",
    [
        "Registrar uma entrada de estoque (recebimento), conforme a sugestão do G4 no enunciado.",
        "O caminho da requisição passa por todas as camadas da Onion:",
    ],
    bullets=[
        "POST /api/recebimentos chega no RecebimentoController (presentation).",
        "O controller chama o caso de uso ProcessarEntrada (application).",
        "O caso de uso usa as entidades Produto, EstoqueItem e Movimentacao do domínio.",
        "Os repositórios da infra (JsonFileEstoqueRepository e JsonFileMovimentacaoRepository) gravam em data/wms-db.json.",
        "Volta um HTTP 201 com o item de estoque criado.",
    ],
)

add_qa(
    "Quais são as principais pastas do projeto?",
    [
        "O código fica em src/, dividido nas quatro camadas da Onion:",
    ],
    bullets=[
        "src/domain — regras de negócio puras, sem dependências externas (entities, enums, repositories como interfaces, services com a PoliticaFifo e errors).",
        "src/application — casos de uso que orquestram o domínio (use-cases de produtos, recebimento, armazenagem, transferência, expedição, estoque e rastreabilidade).",
        "src/infrastructure — implementações técnicas: persistence/JsonDatabase.ts e os JsonFile*Repository que implementam as interfaces do domínio.",
        "src/presentation — API Express, controllers e injeção de dependência em container.ts.",
        "tests/ — teste da regra FIFO isolado (sem banco e sem HTTP).",
        "data/ — arquivo wms-db.json que serve como “banco” entre execuções.",
    ],
)

add_qa(
    "O que está funcionando, mesmo que de forma simples?",
    [
        "A API sobe em http://localhost:3333 e persiste tudo num arquivo JSON entre execuções. "
        "Estão funcionando, de ponta a ponta:",
    ],
    bullets=[
        "Cadastro, atualização e consulta de produtos.",
        "Recebimento — o fluxo escolhido para o Sprint 1.",
        "Armazenagem.",
        "Transferência entre localizações.",
        "Expedição com FIFO.",
        "Consulta de saldo por produto.",
        "Rastreabilidade (histórico de movimentações).",
        "Teste automatizado da regra FIFO em tests/domain/PoliticaFifo.test.ts.",
    ],
)

p = doc.add_paragraph()
run = p.add_run("Como rodar")
run.bold = True
doc.add_paragraph("npm install")
doc.add_paragraph("npm run dev")
doc.add_paragraph("npm test")
doc.add_paragraph(
    "Ao subir a API, o console imprime IDs de usuários e localizações já criados pelo seed — "
    "esses IDs são usados no corpo das requisições."
)

output = r"c:\Users\vitto\Área de Trabalho\Arquitetura - Trabalho WMS\TP2-Sprint1-Respostas.docx"
doc.save(output)
print(f"Arquivo gerado: {output}")
